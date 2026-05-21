"""
生成 SmartPaper-Tagging 系統架構與創新策略 Word 報告
執行：python generate_report.py
輸出：SmartPaper_系統架構與創新策略.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = "SmartPaper_系統架構與創新策略.docx"

# ── 色彩定義 ─────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1E, 0x29, 0x3B)   # 深藍黑
C_BLUE   = RGBColor(0x1D, 0x4E, 0xD8)   # 主藍
C_TEAL   = RGBColor(0x0D, 0x94, 0x88)   # 青綠
C_VIOLET = RGBColor(0x7C, 0x3A, 0xED)   # 紫
C_GREY   = RGBColor(0x64, 0x74, 0x8B)   # 灰
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER_BG = "1E293B"                  # 表頭背景（hex str for shading）
C_ROW_BG    = "EFF6FF"                  # 表格偶數行背景


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFDBFE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or (C_BLUE if level == 1 else C_NAVY)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, size=11, color=None, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_NAVY
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or C_NAVY
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表頭
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, C_HEADER_BG)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = C_WHITE
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 資料列
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = C_ROW_BG if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = cell_text
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.color.rgb = C_NAVY
            set_cell_borders(cell)

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_divider(doc, color="3B82F6"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


# ── 文件主體 ─────────────────────────────────────────────────────────────

doc = Document()

# 頁面邊距
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# 預設字型
doc.styles["Normal"].font.name = "微軟正黑體"
doc.styles["Normal"].font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("SmartPaper-Tagging")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = C_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("智能學術論文管理系統\n系統架構與創新策略")
run.font.size = Pt(16)
run.font.color.rgb = C_NAVY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.date.today().strftime("%Y 年 %m 月 %d 日"))
run.font.size = Pt(12)
run.font.color.rgb = C_GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. 專案定位
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "一、專案定位與目標", level=1)
add_divider(doc)

add_para(doc,
    "SmartPaper-Tagging 是一套以 AI 為核心的學術論文管理平台，"
    "目標是讓研究者從「收集論文」到「寫出論文」的全流程都能在同一個工具內完成，"
    "免去在多個軟體之間切換的摩擦。",
    size=11)

add_para(doc, "核心價值主張：", bold=True, size=11)
bullets = [
    "自動化 → 匯入即標籤、上傳 PDF 即解析，最大化縮短人工整理時間",
    "深度理解 → 不只儲存論文，更理解論文的章節結構、核心概念與引用關係",
    "寫作導向 → 系統終點是幫助使用者「寫出」論文，而非只是「找到」論文",
    "本地優先 → 資料儲存在本機（SQLite + ChromaDB），不依賴雲端訂閱",
]
for b in bullets:
    add_bullet(doc, b)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 2. 系統架構
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "二、系統架構", level=1)
add_divider(doc)

add_para(doc,
    "系統採用「資料層 → 服務層 → UI 層」三層架構，"
    "各層之間透過明確的介面解耦，方便未來擴展或替換底層模型。",
    size=11)

add_heading(doc, "2.1 技術棧", level=2, color=C_TEAL)
add_table(doc,
    headers=["分類", "技術", "用途"],
    rows=[
        ["語言",       "Python 3.11+",                    "全端開發"],
        ["桌面 UI",    "Flet 0.24",                       "跨平台桌面 / Web 介面"],
        ["LLM",        "Google Gemini (google-genai)",    "標籤生成、RAG 問答、缺口分析"],
        ["向量搜尋",   "ChromaDB + allenai-specter",      "語意相似度搜尋"],
        ["關鍵字搜尋", "BM25 (rank-bm25)",                "全文關鍵字搜尋"],
        ["重排序",     "CrossEncoder ms-marco-MiniLM",    "候選論文精確重排"],
        ["對話記憶",   "paraphrase-multilingual-MiniLM",  "跨輪問答記憶萃取"],
        ["PDF 解析",   "pymupdf4llm + pdfplumber",        "Markdown 結構化解析（含 fallback）"],
        ["外部文獻",   "arXiv API",                       "即時查詢外部論文（免費、無 Key）"],
        ["資料庫",     "SQLite + ChromaDB",               "Metadata 與向量持久化儲存"],
    ],
    col_widths=[2.8, 5.0, 5.5])

add_heading(doc, "2.2 模組結構", level=2, color=C_TEAL)
add_table(doc,
    headers=["模組", "主要職責"],
    rows=[
        ["api/",             "Crossref、Gemini、arXiv、Semantic Scholar API 封裝"],
        ["database/",        "SQLite（LRU 快取）、ChromaDB（singleton embedding）、ChunkStore"],
        ["processing/",      "PDF 解析（pymupdf4llm 優先）、HTML 清理、標籤邏輯"],
        ["services/search",  "向量 + BM25 混合搜尋、RRF 融合、LRU query 快取"],
        ["services/reranker","CrossEncoder 重排（BM25 pre-filter + score memoization）"],
        ["services/qa_*",    "多輪 RAG 問答、對話記憶萃取、session 管理"],
        ["services/writing_guide", "三步驟寫作導引、缺口分析、arXiv 外部建議"],
        ["services/classifier",    "語意 / 兩階段 RAG / LLM 三種論文分類"],
        ["ui/views/",        "Flet 各功能介面（QA、寫作導引、知識圖譜⋯）"],
    ],
    col_widths=[4.5, 8.8])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. 核心創新點
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "三、核心創新點", level=1)
add_divider(doc)

add_para(doc,
    "以下七項是本系統有別於市場上現有文獻管理工具（Zotero、Mendeley、Notion）的差異化設計，"
    "每一項都對應一個具體的研究者痛點。",
    size=11)

# ── 創新 1 ──────────────────────────────────────────────────────
add_heading(doc, "創新 1：三步驟寫作引用導引（Writing Guide）", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：研究者在撰寫論文時，常不確定「哪段落應該引用哪篇文獻的哪個概念」，"
    "需要反覆回頭翻閱大量論文。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案：", bold=True)
add_bullet(doc, "Step 1 — 使用者輸入大綱段落，系統以向量搜尋 + CrossEncoder 重排，找出最相關候選論文")
add_bullet(doc, "Step 2 — Gemini 分析每篇候選論文，給出「引用時機」「引用概念」「段落位置（開頭/中間/結尾）」")
add_bullet(doc, "Step 3 — AI 分析整份大綱的概念缺口，查詢本地文獻庫補強，無對應論文時自動查詢 arXiv 建議外部文獻，並附具體寫作範例句")
add_para(doc,
    "差異化：現有工具（如 Connected Papers、Semantic Scholar）只做「找論文」，"
    "本系統做到「告訴你這篇論文在你的文章哪裡用、怎麼用」，並從 arXiv 即時補缺。",
    color=C_BLUE, size=10.5)

doc.add_paragraph()

# ── 創新 2 ──────────────────────────────────────────────────────
add_heading(doc, "創新 2：多路徑 RAG 問答（Triple-Path Retrieval）", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：只用摘要做 RAG 回答粗淺；只用全文又容易漏掉沒上傳 PDF 的論文。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案 — 三條平行檢索路徑：", bold=True)
add_bullet(doc, "路徑 A：全文 chunk 搜尋（有 PDF 的論文，章節感知切割）")
add_bullet(doc, "路徑 B：摘要向量搜尋（補充未上傳 PDF 的論文）")
add_bullet(doc, "路徑 C：追問記憶注入（偵測到追問語意時，強制帶入上一輪引用的論文）")
add_bullet(doc, "CrossEncoder 重排時，Methodology / Results 章節額外加權（重要性係數 × rerank_score）")
add_para(doc,
    "效果：相較於單路徑 RAG，多路徑設計讓回答覆蓋率明顯提升，"
    "尤其在論文庫新舊混雜（部分有 PDF、部分只有摘要）的場景。",
    color=C_BLUE, size=10.5)

doc.add_paragraph()

# ── 創新 3 ──────────────────────────────────────────────────────
add_heading(doc, "創新 3：多 Session 對話紀錄管理", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：傳統 AI 問答介面每次刷新就清空，研究者無法同時維護多個主題的對話脈絡。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案：", bold=True)
add_bullet(doc, "最多保留 5 個獨立 session，每個 session 自動命名（取自第一個問題前 22 字）")
add_bullet(doc, "切換 session 時完整還原對話歷史、來源論文選擇與上一輪結果")
add_bullet(doc, "每個 session 獨立維護對話記憶（ConversationMemory），不互相干擾")
add_bullet(doc, "每次回答後自動提供 2 個「引導式追問」建議，降低使用者思考成本")

doc.add_paragraph()

# ── 創新 4 ──────────────────────────────────────────────────────
add_heading(doc, "創新 4：時間衰減對話記憶系統", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：長對話後，早期提到的重要概念會被遺忘，導致回答前後矛盾。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案：", bold=True)
add_bullet(doc, "對話記憶以語意向量儲存，每個 turn 都有時間戳")
add_bullet(doc, "檢索時使用「語意相似度 × 時間衰減係數」綜合排序，近期記憶優先但不排除舊記憶")
add_bullet(doc, "背景執行萃取（非同步，不阻塞 UI），使用 paraphrase-multilingual-MiniLM 做中英文記憶向量")
add_bullet(doc, "記憶以結構化 JSON 存入，可跨 session 注入（Shared Memory 模式）")

doc.add_paragraph()

# ── 創新 5 ──────────────────────────────────────────────────────
add_heading(doc, "創新 5：章節感知 PDF 解析（Section-Aware Chunking）", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：傳統 PDF 解析將全文切成等長 chunk，方法論和致謝章節的文字被一視同仁。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案：", bold=True)
add_bullet(doc, "pymupdf4llm 將 PDF 轉為 Markdown，利用 ## 標題自動偵測章節邊界")
add_bullet(doc, "每個 chunk 標記所屬章節類型：introduction / methodology / results / discussion / other")
add_bullet(doc, "建立「章節重要性係數」：Methodology = 1.3、Results = 1.2、其他 = 1.0")
add_bullet(doc, "重排序時以 rerank_score × importance_weight 計算最終分數，讓方法論優先出現在回答中")
add_bullet(doc, "表格獨立成 chunk，標記 is_table=True，問答時可引用具體數字")

doc.add_paragraph()

# ── 創新 6 ──────────────────────────────────────────────────────
add_heading(doc, "創新 6：arXiv 即時外部論文建議", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：文獻庫再大，研究者永遠有「這個概念找不到對應論文」的時刻。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案：", bold=True)
add_bullet(doc, "寫作導引 Step 3 發現概念缺口且本地無對應論文時，自動以缺口概念 + 原因查詢 arXiv API")
add_bullet(doc, "結果直接嵌入缺口卡片，顯示標題、作者、摘要預覽")
add_bullet(doc, "「加入文獻庫」一鍵將 arXiv 論文寫入 SQLite + ChromaDB，立即參與後續問答與搜尋")
add_bullet(doc, "免費、無需 API Key，使用 arXiv 官方 Atom feed")

doc.add_paragraph()

# ── 創新 7 ──────────────────────────────────────────────────────
add_heading(doc, "創新 7：純演算法搜尋加速（不換模型，不改架構）", level=2, color=C_VIOLET)
add_para(doc,
    "痛點：本地 NLP 模型（440 MB embedding + CrossEncoder）導致搜尋體感遲緩。",
    italic=True, color=C_GREY)
add_para(doc, "解決方案 — 五項純演算法 / 資料結構優化：", bold=True)
add_table(doc,
    headers=["優化項目", "原理", "效益"],
    rows=[
        ["Embedding 模型 Singleton",  "全 process 共用一個 model 實例",              "消除每次 VectorDB 初始化的 440 MB 重載"],
        ["Paper LRU 記憶體快取",       "OrderedDict（max 2000）+ 執行緒鎖",          "重複論文查詢 0 SQL roundtrip"],
        ["get_by_ids() 批次查詢",      "單一 WHERE IN 取代 N 次 get_by_id()",        "SQL 層加速 5–10×"],
        ["BM25 Pre-filter",           "CrossEncoder 前先以 BM25 縮減候選至 20 篇",  "CrossEncoder 推論量減少 ~60%"],
        ["CrossEncoder Score 快取",   "(query, md5(text)) → score dict",            "相同組合跳過神經網路推論"],
        ["Hybrid Search LRU 快取",    "OrderedDict（max 128 query）",               "相同問題第 2 次起即時回傳"],
    ],
    col_widths=[3.8, 5.0, 4.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. 搜尋與問答流程
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "四、搜尋與問答完整流程", level=1)
add_divider(doc)

add_heading(doc, "4.1 混合搜尋流程", level=2, color=C_TEAL)
steps = [
    ("向量搜尋",    "ChromaDB ANN 搜尋，取前 30 個候選（allenai-specter embedding）"),
    ("BM25 搜尋",   "BM25Okapi 全文搜尋，取前 30 個候選（標題 + 摘要）"),
    ("RRF 融合",    "Reciprocal Rank Fusion 合併兩路排名列表"),
    ("BM25 Pre-filter", "若候選 > 20 篇，先以 BM25 縮減，減少 CrossEncoder 負擔"),
    ("CrossEncoder 重排", "ms-marco-MiniLM 精確計算 query–document 相關性，取前 N 篇"),
    ("快取回傳",    "結果寫入 LRU query 快取，相同查詢下次直接命中"),
]
for i, (title, desc) in enumerate(steps, 1):
    add_para(doc, f"  {i}. {title}：{desc}", size=10.5)

doc.add_paragraph()
add_heading(doc, "4.2 RAG 問答流程", level=2, color=C_TEAL)
steps2 = [
    ("來源選擇", "使用者勾選參與問答的論文（預設全選）"),
    ("三路檢索", "全文 chunks + 摘要向量 + 追問記憶注入（見創新 2）"),
    ("重要性加權重排", "rerank_score × section_importance_weight"),
    ("Context 組裝", "取前 N 個 SourceChunk，格式化為「[編號] 標題 / 章節 / 頁碼」"),
    ("對話記憶注入", "以語意 + 時間衰減找出最相關的歷史記憶，追加到 prompt"),
    ("Gemini 生成", "帶引用編號的學術風格回答，引用來源即時標注"),
    ("追問建議", "獨立 thread 非同步生成 2 個延伸問題 chip"),
    ("記憶萃取", "背景執行，將本輪 Q&A 萃取為結構化記憶存入向量記憶庫"),
]
for i, (title, desc) in enumerate(steps2, 1):
    add_para(doc, f"  {i}. {title}：{desc}", size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. 功能全覽
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "五、功能全覽", level=1)
add_divider(doc)

add_table(doc,
    headers=["功能模組", "主要操作", "技術亮點"],
    rows=[
        ["論文匯入",     "Excel / DOI / arXiv ID 批次匯入",          "Crossref + arXiv 雙源自動補摘要"],
        ["AI 自動標籤",  "Gemini 分析摘要生成標籤",                  "結構化 JSON 輸出，錯誤容錯"],
        ["PDF 全文解析", "上傳 PDF → 章節感知 chunking",             "pymupdf4llm Markdown + importance weight"],
        ["混合搜尋",     "自然語言 / 關鍵字 / 標籤搜尋",             "向量 + BM25 + RRF + CrossEncoder"],
        ["RAG 問答",     "多輪對話問答，最多 5 個 session",          "三路檢索 + 時間衰減記憶 + 追問建議"],
        ["寫作導引",     "輸入大綱，三步驟引用建議",                  "arXiv 外部補缺 + 具體寫作範例"],
        ["論文分類",     "語意 / 兩階段 RAG / LLM 三種方法",        "可批次生成分類總結"],
        ["知識圖譜",     "概念關係視覺化",                           "pyvis 互動圖表"],
        ["文獻分析",     "生成文獻回顧比較表",                       "Gemini 結構化輸出"],
        ["標籤管理",     "重命名 / 合併 / 刪除標籤",                 "SQLite JSON 陣列原地更新"],
        ["引用格式",     "APA / MLA 一鍵複製",                       "自動格式化，支援多作者"],
    ],
    col_widths=[3.2, 5.0, 5.1])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. 與現有工具比較
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "六、與現有工具比較", level=1)
add_divider(doc)

add_table(doc,
    headers=["功能", "Zotero", "Mendeley", "Notion AI", "SmartPaper-Tagging"],
    rows=[
        ["AI 自動標籤",      "✗", "✗", "部分", "✓ Gemini 結構化"],
        ["PDF 全文 RAG 問答", "✗", "✗", "✗",   "✓ 章節級引用"],
        ["多 Session 對話",  "✗", "✗", "✗",   "✓ 最多 5 個"],
        ["寫作引用導引",     "✗", "✗", "✗",   "✓ 三步驟 + 範例句"],
        ["arXiv 外部補缺",   "手動", "手動", "✗", "✓ 自動查詢一鍵匯入"],
        ["混合語意搜尋",     "✗", "✗", "部分", "✓ 向量+BM25+重排"],
        ["對話記憶",         "✗", "✗", "✗",   "✓ 時間衰減向量記憶"],
        ["本地部署",         "✓", "✗", "✗",   "✓ 完全離線可用"],
        ["免費使用",         "✓", "部分", "✗", "✓ 只需 Gemini Key"],
    ],
    col_widths=[4.2, 2.0, 2.2, 2.5, 3.8])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 7. 未來方向
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "七、未來可擴展方向", level=1)
add_divider(doc)

future = [
    ("多用戶 / 團隊協作", "目前單用戶本地部署；可加入 FastAPI 後端 + JWT 認證，支援實驗室共享文獻庫"),
    ("GPU 加速推論",       "將 embedding 與 CrossEncoder 移至 GPU（CUDA），搜尋速度可再提升 5–10×"),
    ("Semantic Scholar 引用圖", "目前已整合引用資料擷取，可進一步實作「引用鏈追蹤」與影響力排名"),
    ("論文推薦系統",       "基於閱讀歷史與標籤偏好，主動推薦相關論文"),
    ("Gemini Function Calling", "讓 LLM 自主決定何時搜尋 / 查外部 API，實現更靈活的 Agent 模式"),
    ("Web 版本",           "Flet 已支援 Web 模式，可部署為研究機構的論文管理入口"),
]
for title, desc in future:
    add_para(doc, f"▸ {title}", bold=True, size=11, color=C_BLUE)
    add_para(doc, f"  {desc}", size=10.5, color=C_NAVY, indent=0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════════
# 頁尾
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_p.add_run("— 文件結束 —")
run.font.color.rgb = C_GREY
run.font.size = Pt(11)
run.font.italic = True

doc.save(OUTPUT)
print(f"已生成：{OUTPUT}")
